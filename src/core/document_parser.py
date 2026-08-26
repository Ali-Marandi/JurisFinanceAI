"""
JurisFinanceAI - Document Parser
Handles parsing of PDF, DOCX, and text files.
"""

import os
from pathlib import Path
from typing import Optional, Dict, Tuple


class DocumentParser:
    """Parses various document formats for analysis."""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'text',
        '.rtf': 'text',
    }

    @staticmethod
    def parse(filepath: str) -> Tuple[Optional[str], Optional[Dict]]:
        """
        Parse a document and return its text content and metadata.

        Args:
            filepath: Path to the document file

        Returns:
            Tuple of (text_content, metadata_dict) or (None, None) on failure
        """
        path = Path(filepath)
        if not path.exists():
            return None, None

        ext = path.suffix.lower()
        file_type = DocumentParser.SUPPORTED_EXTENSIONS.get(ext)

        if file_type == 'pdf':
            return DocumentParser._parse_pdf(filepath)
        elif file_type == 'docx':
            return DocumentParser._parse_docx(filepath)
        elif file_type == 'text':
            return DocumentParser._parse_text(filepath)
        else:
            return None, {"error": f"Unsupported file type: {ext}"}

    @staticmethod
    def _parse_pdf(filepath: str) -> Tuple[Optional[str], Dict]:
        """Parse a PDF file."""
        try:
            import pdfplumber
            text_parts = []
            metadata = {
                "page_count": 0,
                "file_size": os.path.getsize(filepath),
            }

            with pdfplumber.open(filepath) as pdf:
                metadata["page_count"] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- صفحه {i + 1} ---\n{page_text}")

                # Try to extract tables
                tables = []
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            tables.append(table)

                if tables:
                    metadata["has_tables"] = True
                    metadata["table_count"] = len(tables)

            full_text = "\n\n".join(text_parts)
            return full_text, metadata

        except Exception as e:
            return None, {"error": str(e)}

    @staticmethod
    def _parse_docx(filepath: str) -> Tuple[Optional[str], Dict]:
        """Parse a DOCX file."""
        try:
            from docx import Document
            doc = Document(filepath)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(tables),
                "file_size": os.path.getsize(filepath),
                "has_tables": len(tables) > 0,
            }

            full_text = "\n".join(text_parts)
            return full_text, metadata

        except Exception as e:
            return None, {"error": str(e)}

    @staticmethod
    def _parse_text(filepath: str) -> Tuple[Optional[str], Dict]:
        """Parse a plain text file."""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()

            metadata = {
                "char_count": len(text),
                "word_count": len(text.split()),
                "line_count": text.count("\n") + 1,
                "file_size": os.path.getsize(filepath),
            }

            return text, metadata

        except Exception as e:
            return None, {"error": str(e)}

    @staticmethod
    def get_file_info(filepath: str) -> Dict:
        """Get basic file information."""
        path = Path(filepath)
        if not path.exists():
            return {"error": "File not found"}

        stat = path.stat()
        return {
            "name": path.name,
            "path": str(path.absolute()),
            "extension": path.suffix.lower(),
            "size_bytes": stat.st_size,
            "size_mb": round(stat.st_size / (1024 * 1024), 2),
            "modified": stat.st_mtime,
        }

    @staticmethod
    def is_supported(filepath: str) -> bool:
        """Check if a file format is supported."""
        ext = Path(filepath).suffix.lower()
        return ext in DocumentParser.SUPPORTED_EXTENSIONS
